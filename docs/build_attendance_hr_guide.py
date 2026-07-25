from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path

OUT = Path(__file__).with_name("HRMS_Attendance_Rules_HR_Guide_Updated.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
MUTED = "666666"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GOLD = "D49A16"
WHITE = "FFFFFF"
INK = "222222"


def set_font(run, name="Calibri", size=11, bold=None, italic=None, color=INK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Callout"]
    callout.font.name = "Calibri"
    callout.font.size = Pt(10.5)
    callout.font.color.rgb = RGBColor.from_string(NAVY)
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.18)
    callout.paragraph_format.space_before = Pt(5)
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.line_spacing = 1.2


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.62)
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE):
    p = doc.add_paragraph()
    p.style = doc.styles["Callout"]
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "16")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), BLUE)
    borders.append(left)
    p_pr.append(borders)
    r = p.add_run(f"{label}: ")
    set_font(r, size=10.5, bold=True, color=NAVY)
    r = p.add_run(text)
    set_font(r, size=10.5, color=NAVY)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        set_cell_shading(hdr.cells[i], LIGHT_BLUE)
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(header)
        set_font(r, size=9.5, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(str(value))
            set_font(r, size=9.5, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)
configure_styles(doc)

# Running header/footer.
hp = section.header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hr = hp.add_run("HRMS Attendance Policy & Setup Guide")
set_font(hr, size=9, bold=True, color=MUTED)
add_page_number(section.footer.paragraphs[0])

# Editorial cover.
doc.add_paragraph().paragraph_format.space_after = Pt(54)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("HR OPERATIONS GUIDE")
set_font(r, size=11, bold=True, color=GOLD)
p.paragraph_format.space_after = Pt(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Employee Attendance Rules")
set_font(r, size=28, bold=True, color=NAVY)
p.paragraph_format.space_after = Pt(5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Shift Setup, Biometric Processing, Exceptions & Permissions")
set_font(r, size=14, color=DARK_BLUE)
p.paragraph_format.space_after = Pt(26)

add_callout(
    doc,
    "Purpose",
    "Yeh guide HR team ko current HRMS attendance rules samjhane, shifts configure karne, biometric mapping verify karne aur attendance exceptions handle karne ke liye tayar ki gayi hai.",
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(55)
r = p.add_run("Prepared for HR Department")
set_font(r, size=12, bold=True, color=NAVY)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Current system policy reference | July 2026")
set_font(r, size=10, italic=True, color=MUTED)

doc.add_page_break()

doc.add_heading("1. Policy at a Glance", level=1)
add_callout(
    doc,
    "Core principle",
    "Attendance department ke naam se directly calculate nahi hoti. Employee ki assigned shift, working days, confirmed holiday/closure, approved leave aur corrected biometric punch time final result decide karte hain.",
)

add_table(
    doc,
    ["Shift category", "Grace", "Arrival half-day", "Full duty", "Half-day work"],
    [
        ["> 7-hour fixed", "15 min", "> 150 min late", "Assigned required minutes", "50% required minutes"],
        ["<= 7-hour fixed", "0 min", "> 120 min late", "Assigned required minutes", "50% required minutes"],
        ["Flexible 8-hour", "Not applicable", "No late status", "480 min", "240 min"],
        ["Flexible 6-hour", "Not applicable", "No late status", "360 min", "180 min"],
        ["Saturday scheduled", "No late", "No half-day", "Any valid sign-in = Present", "No half-day"],
    ],
    [1900, 1200, 1900, 2260, 2100],
)

doc.add_heading("2. What Controls Attendance?", level=1)
for item in [
    "Assigned shift and whether that shift is active.",
    "Shift type: Fixed or Flexible.",
    "Shift start/end time and whether it crosses midnight.",
    "Required minutes and working days.",
    "Employee department when a holiday is department-specific.",
    "Confirmed holiday, half-day, early closure or late opening.",
    "Approved leave status.",
    "Biometric machine User ID mapping and corrected timestamp.",
]:
    add_bullet(doc, item)

doc.add_heading("3. Fixed Shift Rules", level=1)
doc.add_heading("3.1 Fixed shift longer than seven hours", level=2)
doc.add_paragraph("Example: 8:00 PM to 4:00 AM (8 hours).")
add_table(
    doc,
    ["Arrival", "Result"],
    [
        ["8:00 PM - 8:15 PM", "Present"],
        ["8:16 PM", "Late"],
        ["10:30 PM", "Late"],
        ["10:31 PM onward", "Half Day"],
    ],
    [3100, 6260],
)
add_callout(
    doc,
    "Important",
    "Arrival half-day boundary cross karne ke baad employee raw duty hours baad mein complete kar le, tab bhi attendance Half Day hi rehti hai.",
    LIGHT_GRAY,
)

doc.add_heading("3.2 Fixed shift of seven hours or less", level=2)
doc.add_paragraph("Example: 10:00 AM to 4:30 PM (6 hours 30 minutes).")
add_table(
    doc,
    ["Arrival", "Result"],
    [
        ["10:00 AM", "Present"],
        ["10:01 AM", "Late"],
        ["12:00 PM", "Late"],
        ["12:01 PM onward", "Half Day"],
    ],
    [3100, 6260],
)
doc.add_paragraph("Seven hours ya kam fixed shift ko grace period nahi milta.")

doc.add_heading("4. Flexible Shift Rules", level=1)
doc.add_paragraph("Flexible shift mein employee kisi bhi time sign in kar sakta hai. Late aur grace calculation apply nahi hoti.")
add_table(
    doc,
    ["Flexible type", "Present", "Half Day", "Absent"],
    [
        ["8-hour", "8 hours complete", "4 to <8 hours", "<4 hours"],
        ["6-hour", "6 hours complete", "3 to <6 hours", "<3 hours"],
    ],
    [1900, 2500, 2500, 2460],
)
add_callout(doc, "Configuration limit", "Flexible shifts sirf 6-hour ya 8-hour required duration ke saath create ho sakti hain.", LIGHT_GRAY)

doc.add_heading("5. Worked Hours and Final Status", level=1)
doc.add_paragraph("Sign-out ke waqt system yeh formula use karta hai:")
add_callout(doc, "Formula", "Worked Minutes = Sign Out - Sign In. Break-time deduction apply nahi hoti.")
for item in [
    "Required minutes complete: Present; agar arrival late thi to Late.",
    "Half-day minutes complete magar full requirement se kam: Half Day.",
    "Half-day minutes se kam: Absent.",
    "Fixed shift employee arrival half-day threshold cross kare to status Half Day hi rehta hai.",
]:
    add_bullet(doc, item)

doc.add_heading("6. Overnight Shift", level=1)
doc.add_paragraph("Example: 6:00 PM to 2:00 AM.")
add_table(
    doc,
    ["Event", "Timestamp", "Duty date"],
    [
        ["Sign In", "26 July, 6:00 PM", "26 July"],
        ["Sign Out", "27 July, 2:00 AM", "26 July"],
    ],
    [1800, 3500, 4060],
)
doc.add_paragraph("Attendance us calendar date ke under rakhi jati hai jis din shift start hui. After-midnight punch previous overnight shift se match ho sakta hai.")

doc.add_heading("7. Biometric Attendance", level=1)
doc.add_heading("7.1 Timestamp correction", level=2)
doc.add_paragraph("Machine ki clock business requirement ke mutabiq 12 hours behind rakhi gayi hai. HRMS machine setting change nahi karta.")
add_callout(doc, "Current environment", "BIOMETRIC_TIME_OFFSET_HOURS=12")
add_table(
    doc,
    ["Machine timestamp", "Corrected system timestamp"],
    [
        ["26 July, 6:00 AM", "26 July, 6:00 PM"],
        ["26 July, 2:00 PM", "27 July, 2:00 AM"],
        ["31 July, 11:30 PM", "1 August, 11:30 AM"],
    ],
    [4300, 5060],
)
doc.add_paragraph("Raw machine timestamp audit ke liye preserve hoti hai. Shift, late, half-day, payroll, reports aur dashboard sirf corrected timestamp use karte hain.")

doc.add_heading("7.2 Biometric punch sequence", level=2)
for text in [
    "First valid punch for the shift -> Sign In.",
    "Second valid punch for the same open shift -> Sign Out.",
    "Third or extra punch -> Ignored.",
    "Duplicate device event -> Ignored through fingerprint deduplication.",
]:
    add_number(doc, text)
add_callout(
    doc,
    "Operational note",
    "Machine par IN/OUT selection currently final action decide nahi karti. HRMS existing open shift record dekh kar first/second punch sequence decide karta hai.",
    LIGHT_GRAY,
)

doc.add_heading("7.3 Employee mapping", level=2)
for text in [
    "Machine mein employee ka actual User ID note karein.",
    "HRMS > Employees > Edit Employee open karein.",
    "Biometric Device User ID mein exact same value enter karein.",
    "Save karke fresh punch test karein.",
]:
    add_number(doc, text)
doc.add_paragraph("Example: Machine User ID 25 = HRMS Biometric Device User ID 25. Duplicate mapping allowed nahi.")

doc.add_heading("8. Manual Sign In and Sign Out", level=1)
doc.add_paragraph("Fixed shift manual sign-in window shift start se 4 hours pehle se shift end ke 4 hours baad tak hoti hai. Flexible shift is restriction se exempt hai.")
for item in [
    "Biometric punch manual window restriction bypass karta hai.",
    "Confirmed full-day holiday restriction biometric par bhi apply hoti hai.",
    "Non-working day par normal employee ka sign-in reject hota hai.",
    "Inactive assigned shift par attendance reject hoti hai.",
]:
    add_bullet(doc, item)

doc.add_heading("9. Working Days and Weekly Off", level=1)
add_table(
    doc,
    ["Number", "Day", "Number", "Day"],
    [
        ["0", "Sunday", "4", "Thursday"],
        ["1", "Monday", "5", "Friday"],
        ["2", "Tuesday", "6", "Saturday"],
        ["3", "Wednesday", "", ""],
    ],
    [1100, 3580, 1100, 3580],
)
doc.add_paragraph("Employee ki assigned shift ke workingDays determine karte hain ke us din attendance required hai ya weekly off.")

doc.add_heading("9.1 Saturday Special Rule", level=2)
add_callout(
    doc,
    "Saturday policy",
    "Saturday ko normal attendance status sirf Present ya Absent hoga. Late aur Half Day apply nahi honge.",
)
for item in [
    "Employee ki assigned shift mein Saturday working day ho aur valid sign-in aaye: Present.",
    "Late arrival: phir bhi Present; lateMinutes zero.",
    "Required hours se kam work: phir bhi Present; Half Day nahi.",
    "Sign-out miss ho: attendance Present rahegi; missed sign-out late penalty nahi.",
    "Scheduled Saturday par sign-in bilkul na ho: Absent.",
    "Agar assigned shift mein Saturday working day nahi: weekly off; Absent nahi.",
    "Approved leave ya confirmed holiday apna respective status preserve karega.",
]:
    add_bullet(doc, item)

doc.add_heading("10. Department-wise Setup", level=1)
add_callout(
    doc,
    "Key rule",
    "Department timing automatically define nahi karta. Different department rules ke liye separate shifts create karke relevant employees ko assign karna zaroori hai.",
)
add_table(
    doc,
    ["Department", "Suggested shift", "Example timing", "Assignment"],
    [
        ["Operations", "Operations Night", "8:00 PM - 4:00 AM", "Operations employees"],
        ["Sales", "Sales Day", "10:00 AM - 4:30 PM", "Sales employees"],
        ["Flexible team", "Flexible 6h/8h", "Any start time", "Selected employees"],
    ],
    [1700, 2500, 2200, 2960],
)

doc.add_heading("11. Shift Setup by HR", level=1)
doc.add_paragraph("Shifts sirf HR ya Super Admin create, edit aur delete kar sakte hain.")
for text in [
    "Shift name aur unique code enter karein.",
    "Fixed ya Flexible type select karein.",
    "Fixed shift ke liye start time aur end time configure karein; required duty poori shift window hoti hai.",
    "Working days select karein.",
    "Shift ko active rakhein.",
    "Employee profile mein correct shift assign karein.",
]:
    add_number(doc, text)
doc.add_paragraph("Assigned employees wali shift delete nahi ho sakti; pehle employees ko reassign karna hoga.")

doc.add_heading("12. Employee Without an Assigned Shift", level=1)
doc.add_paragraph("Specific shift na hone par General Shift company settings se use hoti hai.")
for item in [
    "Office start/end: Company Settings.",
    "Weekend days: Company Settings.",
    "Default required duty: 8 hours.",
    "Default half-day: 4 hours.",
]:
    add_bullet(doc, item)
add_callout(doc, "Recommendation", "Har employee ko explicit shift assign karein taake ambiguity na ho.", LIGHT_GRAY)

doc.add_heading("13. Holidays and Office Closures", level=1)
doc.add_paragraph("Holiday sirf HR confirmation ke baad attendance par apply hoti hai.")
add_table(
    doc,
    ["Event type", "Attendance effect"],
    [
        ["Full Day", "Sign-in required nahi; status Holiday."],
        ["Half Day", "Effective required minutes half."],
        ["Early Closure", "Effective end time aur required work reduce."],
        ["Late Opening", "Late calculation new opening time se."],
    ],
    [2400, 6960],
)
doc.add_paragraph("Closure scope All, Department ya Shift ho sakta hai. Department matching case-insensitive hai. Closure paid ya unpaid mark ho sakti hai.")

doc.add_heading("14. Leave, Absence and Daily Reconciliation", level=1)
for item in [
    "Approved leave available: attendance status On Leave.",
    "No approved leave and no attendance on a working day: Absent.",
    "Confirmed full-day holiday: Holiday/off.",
    "Configured weekend: automatic Absent record nahi banta.",
    "Attendance reset date se purane absent records regenerate nahi hote.",
]:
    add_bullet(doc, item)

doc.add_heading("15. Missed Punch Rules", level=1)
doc.add_heading("15.1 Missed Sign In", level=2)
for item in [
    "Previous-day reconciliation Absent record create karti hai.",
    "Missed sign-in ek late violation count karta hai.",
    "Employee ko notification aur regularization option milta hai.",
]:
    add_bullet(doc, item)

doc.add_heading("15.2 Missed Sign Out", level=2)
for item in [
    "System scheduled end ko fake sign-out nahi banata.",
    "Worked hours assume nahi kiye jate.",
    "Status Incomplete aur worked minutes zero hote hain.",
    "Missed sign-out ek late violation count karta hai.",
    "Employee correction request submit kar sakta hai.",
]:
    add_bullet(doc, item)

doc.add_heading("16. Late Count Policy", level=1)
doc.add_paragraph("Har late sign-in aur eligible missed punch employee lateCount increment karta hai.")
add_callout(
    doc,
    "Every third late",
    "Payroll mein har 3 Late ka group 0.5 daily salary deduction banata hai. Agar HR Settings mein per-minute mode select ho to configured per-minute rate use hota hai.",
    LIGHT_GRAY,
)
doc.add_paragraph("Approved late waiver late minutes zero karta hai, Late ko Present bana sakta hai aur employee late count ek kam karta hai.")

doc.add_heading("17. Manual Correction and Regularization", level=1)
doc.add_heading("17.1 HR manual correction", level=2)
for item in [
    "Manual correction sirf HR aur Super Admin kar sakte hain.",
    "Selected attendance ki sign-in date fixed hoti hai; sirf time change hota hai.",
    "Overnight fixed shift ka sign-out next date par fixed hota hai.",
    "Sign-out hamesha sign-in ke baad hona chahiye.",
    "Correction ke baad late, early leave, worked time aur status recalculate hote hain.",
]:
    add_bullet(doc, item)

doc.add_heading("17.2 Employee regularization", level=2)
doc.add_paragraph("Employee sirf apni attendance ke liye Time Correction ya Late Waiver request submit kar sakta hai.")
doc.add_paragraph("Approver resolution current system mein Manager -> Team Lead -> HR -> fallback Super Admin ho sakti hai. Manager/Team Lead sirf assigned request review karte hain; HR elevated reviewer hai.")

doc.add_heading("18. Attendance Visibility and Permissions", level=1)
add_table(
    doc,
    ["Role", "Attendance visibility", "Correction control"],
    [
        ["Employee", "Own attendance only", "Regularization request"],
        ["Team Lead", "Assigned team members", "Assigned request review"],
        ["Manager", "Direct reports", "Assigned request review"],
        ["HR", "All normal company employees", "Manual correction and review"],
        ["Super Admin", "All employees", "Full correction and review"],
    ],
    [1700, 4900, 2760],
)
doc.add_paragraph("Non-Super Admin roles Super Admin attendance nahi dekh sakte.")

doc.add_heading("19. Super Admin Attendance Exemption", level=1)
for item in [
    "Late, Half Day aur Absent penalties apply nahi hoti.",
    "Off-day restriction apply nahi hoti.",
    "Open attendance scheduled end par auto-close ho sakti hai.",
    "Status Present normalize hota hai.",
]:
    add_bullet(doc, item)

doc.add_heading("20. Live Updates, API Sync and Downstream Effects", level=1)
doc.add_paragraph("Active HRMS modules operational data API/database se load karte hain. Successful changes ke baad local cache aur Socket.IO doosre active users ke relevant modules refresh karte hain:")
for item in [
    "Attendance change -> Attendance, Payroll, Dashboard and Reports.",
    "Leave apply/approve/cancel -> Leaves, Employee balance, Attendance, Payroll, Dashboard and Reports.",
    "Employee profile/status/shift change -> Employees, Attendance, Leaves, Payroll, Dashboard, Projects and Reports.",
    "Holiday or company policy change -> Settings/Holidays, Attendance, Payroll, Dashboard and Reports.",
    "Expense or category change -> Expenses, Dashboard and Reports.",
]:
    add_bullet(doc, item)
add_callout(
    doc,
    "Dependency",
    "Biometric live update tabhi aayegi jab device punch backend ko successfully receive ho, employee mapping valid ho aur browser ka Socket.IO session connected ho. Reconnect par application missed data ka catch-up refresh karti hai; notifications ke liye polling fallback bhi active hai.",
    LIGHT_GRAY,
)

doc.add_heading("21. Leave Management Rules", level=1)
doc.add_heading("21.1 Leave types and balances", level=2)
for item in [
    "Company Settings mein enabled leave types hi har role ke Apply Leave form mein dikhte hain.",
    "Casual Leave ko Annual Leave replace kar chuki hai.",
    "Paid balance types Annual aur Sick hain; HR opening available/used balances initialize kar sakta hai.",
    "Leave screen available, used aur remaining balance type-wise show karti hai.",
    "Unpaid leave paid balance consume nahi karti magar payroll deduction create karti hai.",
    "Unused leave next calendar year carry forward nahi hoti; January 1 par balance new annual entitlement par reset hota hai.",
]:
    add_bullet(doc, item)

doc.add_heading("21.2 Date and duty-day calculation", level=2)
for item in [
    "Single-day leave ke liye sirf wohi ek date select karna sufficient hai; separate next-day end date required nahi.",
    "Overnight shift mein 22 July 10:00 PM se 23 July 4:00 AM ka interval 22 July ka ek duty day consume karta hai.",
    "Genuine multiple-date leave inclusive duty dates count karti hai.",
    "Assigned shift working days/weekends leave duty-day calculation mein use hote hain.",
]:
    add_bullet(doc, item)

doc.add_heading("21.3 Three-month eligibility and HR exception", level=2)
for item in [
    "Employee apni joining date ke 3 complete calendar months se pehle leave apply nahi kar sakta.",
    "Is probation period mein sirf HR employee ke existing Absent ya Half Day attendance record par exception leave laga sakta hai.",
    "HR exception ek fixed attendance date aur linked attendance record ke saath hoti hai.",
    "Eligible employee ke liye normal leave workflow use hota hai; probation exception route use nahi hota.",
]:
    add_bullet(doc, item)

doc.add_heading("21.4 Leave approval workflow", level=2)
add_table(
    doc,
    ["Stage", "Reviewer", "Result"],
    [
        ["1", "Assigned Team Lead or Manager", "Approve to HR, or reject"],
        ["2", "HR", "Final approval or rejection"],
        ["No stage-1 assignee", "HR", "Request routes directly to HR"],
    ],
    [1200, 3100, 5060],
)
add_callout(
    doc,
    "Control",
    "Super Admin leave approval stage mein included nahi. Final leave approval HR ke paas hai. Team Lead/Manager sirf apne assigned employees ki request handle karte hain.",
)

doc.add_heading("22. Payroll and Salary Rules", level=1)
add_callout(doc, "Daily salary formula", "Daily Salary = Monthly Salary / 30. Calendar month ke days ya working-day count se divisor change nahi hota.")
add_table(
    doc,
    ["Attendance/pay item", "Salary effect"],
    [
        ["Present / Late (before grouped deduction)", "Full daily salary"],
        ["Half Day", "50% daily salary deduction"],
        ["Absent or Incomplete", "Full daily salary deduction"],
        ["Paid Annual/Sick Leave", "No attendance deduction"],
        ["Unpaid Leave", "Full daily salary deduction"],
        ["Weekly Off / Holiday", "No normal attendance deduction"],
        ["Sandwich Leave", "Full daily salary deduction for each sandwiched off day"],
        ["Bonus / Incentive / Allowance", "Added to gross salary"],
        ["Loan / Advance / Tax", "Deducted with visible reason/breakdown"],
    ],
    [3600, 5760],
)

doc.add_heading("22.1 Sandwich leave policy", level=2)
doc.add_paragraph("Agar weekly off ya confirmed holiday block ke dono sides par Absent ya Unpaid Leave ho to beech ke tamam consecutive off/holiday dates Sandwich Leave ban kar unpaid deduction mein include hoti hain.")
add_callout(doc, "Example", "Saturday Absent + Sunday Weekly Off + Monday Unpaid Leave = Saturday, Sunday aur Monday teen deduction days.")

doc.add_heading("22.2 Late deduction settings", level=2)
for item in [
    "Default mode: 3 Late = 0.5 daily salary deduction.",
    "Multiple complete groups apply ho sakte hain; remainder next complete group tak direct deduction nahi banata.",
    "Optional mode: total late minutes x configured per-minute rate.",
    "Approved late waiver attendance aur payroll late totals ko recalculate karta hai.",
]:
    add_bullet(doc, item)

doc.add_heading("22.3 Payroll workflow and visibility", level=2)
for item in [
    "Payslip workflow: Draft -> Pending Approval -> Approved -> Paid -> Locked.",
    "Approved, Paid ya Locked payslip editable nahi.",
    "Har user apna payroll/payslip dekh sakta hai.",
    "Team Lead aur Manager doosre employees ki salary/payroll details nahi dekh sakte.",
    "Authorized HR/Admin roles company payroll ko role policy ke mutabiq manage karte hain.",
    "System mein overtime payment concept disabled hai; overtime amount net salary mein add nahi hota.",
]:
    add_bullet(doc, item)

doc.add_heading("23. Security and Data Integrity", level=1)
for item in [
    "All operational APIs company/tenant scope enforce karti hain.",
    "Employee sirf apna attendance, leave aur payslip data access karta hai.",
    "Team Lead/Manager access reporting mapping tak limited hai.",
    "Salary controls UI ke saath backend API permissions par bhi protected hain.",
    "Biometric duplicate punches fingerprint deduplication se prevent hote hain.",
    "Raw biometric machine timestamp audit ke liye preserve aur corrected timestamp calculations ke liye use hoti hai.",
    "Form drafts, theme aur notification-sound preference local browser mein ho sakte hain; operational HR records database se aate hain.",
]:
    add_bullet(doc, item)

doc.add_heading("24. Attendance Status Reference", level=1)
add_table(
    doc,
    ["Status", "Meaning"],
    [
        ["Present", "Full required duty complete."],
        ["Late", "Arrival grace ke baad, full duty otherwise complete."],
        ["Half Day", "Arrival threshold cross ya half-day work complete."],
        ["Absent", "No valid attendance ya half-day minimum se kam work."],
        ["Incomplete", "Sign-in hai magar valid sign-out missing."],
        ["On Leave", "Approved leave recorded."],
        ["Holiday", "Confirmed applicable company closure."],
        ["Weekend", "Configured non-working day."],
    ],
    [2200, 7160],
)

doc.add_heading("25. HR Setup Checklist", level=1)
for text in [
    "Company timezone Asia/Karachi verify karein.",
    "Company weekend days configure karein.",
    "Departments verify karein.",
    "Har timing pattern ke liye separate shift create karein.",
    "Shift working days verify karein.",
    "Har employee ko correct active shift assign karein.",
    "Manager aur Team Lead reporting mapping set karein.",
    "Biometric Device User ID exact machine User ID se match karein.",
    "BIOMETRIC_TIME_OFFSET_HOURS=12 verify karein.",
    "Department/shift-specific holidays configure aur confirm karein.",
    "Annual aur Sick opening leave balances initialize karein.",
    "Late deduction mode aur rate verify karein.",
    "Employee monthly salary configure karein; daily salary /30 verify karein.",
    "Test employee ka fresh Sign In aur Sign Out karein.",
    "Attendance, leave balance, dashboard, reports aur payroll recalculation verify karein.",
]:
    add_number(doc, text)

doc.add_heading("26. Physical Biometric Acceptance Test", level=1)
add_table(
    doc,
    ["Step", "Expected result"],
    [
        ["Backend starts", "Device connected and polling/realtime active."],
        ["Mapped employee first punch", "Attendance created; action sign_in."],
        ["Same employee second punch", "Same attendance updated; action sign_out."],
        ["Machine time +12 correction", "Correct date/time including midnight rollover."],
        ["Duplicate event", "Ignored; no duplicate attendance."],
        ["Frontend open", "Attendance appears and live modules refresh."],
    ],
    [3600, 5760],
)

doc.add_heading("27. Final HR Notes", level=1)
add_callout(
    doc,
    "Most important",
    "Department khud attendance timing decide nahi karta. Employee ki assigned shift timing aur working days decide karte hain. Department targeted holiday/closure ke scope ke liye use hota hai.",
)
doc.add_paragraph("Policy ya timing change karte waqt existing open attendance records apne sign-in ke waqt captured shift snapshot par complete hote hain; shift edit open record ki requirement ko retroactively change nahi karti.")

# Prevent lonely headings where practical.
for paragraph in doc.paragraphs:
    if paragraph.style.name.startswith("Heading"):
        paragraph.paragraph_format.keep_with_next = True

doc.core_properties.title = "HRMS Attendance Rules - HR Guide"
doc.core_properties.subject = "Attendance policy, biometric rules, shift setup and permissions"
doc.core_properties.author = "HRMS"
doc.core_properties.keywords = "HRMS, Attendance, HR, Biometric, Shift Policy"

doc.save(OUT)
print(OUT)
